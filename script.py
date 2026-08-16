"""
Resume Text Extractor
----------------------
Extracts text from all PDF and DOCX resumes in a folder and saves each
as a corresponding .txt file (same filename, .txt extension).

Setup:
    pip install pdfplumber python-docx

Usage:
    python extract_resume_text.py --input_dir ./resumes --output_dir ./resume_texts
"""

import argparse
import os
import sys

import pypdfium2 as pdfium
from docx import Document


def extract_text_from_pdf(filepath: str) -> str:
        text = "\n".join(
            p.get_textpage().get_text_range()
            for p in pdfium.PdfDocument(filepath)
        )
        return text


def extract_text_from_docx(filepath: str) -> str:
    doc = Document(filepath)
    text_parts = [para.text for para in doc.paragraphs if para.text.strip()]

    # Also grab text inside tables (some resumes use table layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def process_folder(input_dir: str, output_dir: str) -> None:
    os.makedirs(output_dir, exist_ok=True)

    supported_ext = (".pdf", ".docx")
    files = [f for f in os.listdir(input_dir) if f.lower().endswith(supported_ext)]

    if not files:
        print(f"No PDF or DOCX files found in {input_dir}")
        return

    print(f"Found {len(files)} resume(s) to process.\n")

    success_count = 0
    fail_count = 0

    for filename in files:
        input_path = os.path.join(input_dir, filename)
        base_name = os.path.splitext(filename)[0]
        output_path = os.path.join(output_dir, base_name + ".txt")

        try:
            if filename.lower().endswith(".pdf"):
                text = extract_text_from_pdf(input_path)
            else:  # .docx
                text = extract_text_from_docx(input_path)

            if not text.strip():
                print(f"[WARNING] No extractable text found in: {filename} "
                      f"(may be a scanned/image-based file)")

            with open(output_path, "w", encoding="utf-8") as out_file:
                out_file.write(text)

            print(f"[OK] {filename} -> {os.path.basename(output_path)}")
            success_count += 1

        except Exception as e:
            print(f"[FAILED] {filename}: {e}")
            fail_count += 1

    print(f"\nDone. {success_count} succeeded, {fail_count} failed.")
    print(f"Text files saved to: {os.path.abspath(output_dir)}")


def main():
    parser = argparse.ArgumentParser(description="Extract text from PDF/DOCX resumes into .txt files.")
    parser.add_argument("--input_dir", required=True, help="Folder containing resume PDF/DOCX files")
    parser.add_argument("--output_dir", default="./resume_texts", help="Folder to save extracted .txt files")
    args = parser.parse_args()

    if not os.path.isdir(args.input_dir):
        print(f"Input directory not found: {args.input_dir}")
        sys.exit(1)

    process_folder(args.input_dir, args.output_dir)


if __name__ == "__main__":
    main()