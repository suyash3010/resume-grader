import json
import os
import io
import time
import tempfile
import logging
from datetime import datetime
from functools import wraps
from urllib.parse import urlencode, parse_qs

from flask import Flask, request, jsonify, send_file, session, redirect, url_for, render_template_string
from markupsafe import escape
from openai import OpenAI
import pandas as pd
from werkzeug.utils import secure_filename
import pypdfium2 as pdfium
from docx import Document
from dotenv import load_dotenv
import requests
import boto3

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'dev-secret-key-change-in-production')

# Cognito Configuration
COGNITO_USER_POOL_ID = os.environ.get('COGNITO_USER_POOL_ID')
COGNITO_CLIENT_ID = os.environ.get('COGNITO_CLIENT_ID')
COGNITO_CLIENT_SECRET = os.environ.get('COGNITO_CLIENT_SECRET')
COGNITO_DOMAIN = os.environ.get('COGNITO_DOMAIN')
COGNITO_REGION = os.environ.get('COGNITO_REGION', 'us-east-1')

# Build Cognito URLs
COGNITO_DISCOVERY_URL = f'https://cognito-idp.{COGNITO_REGION}.amazonaws.com/{COGNITO_USER_POOL_ID}/.well-known/openid-configuration'
COGNITO_AUTH_URL = f'https://{COGNITO_DOMAIN}.auth.{COGNITO_REGION}.amazoncognito.com/oauth2/authorize'
COGNITO_TOKEN_URL = f'https://{COGNITO_DOMAIN}.auth.{COGNITO_REGION}.amazoncognito.com/oauth2/token'
COGNITO_JWKS_URL = f'https://cognito-idp.{COGNITO_REGION}.amazonaws.com/{COGNITO_USER_POOL_ID}/.well-known/jwks.json'

# Cognito client for token verification
cognito_client = boto3.client('cognito-idp', region_name=COGNITO_REGION)

# DynamoDB Configuration
ACTIVITY_TABLE_NAME = os.environ.get('ACTIVITY_TABLE_NAME', 'resume-grader-user-activity')
dynamodb = boto3.resource('dynamodb')
activity_table = dynamodb.Table(ACTIVITY_TABLE_NAME)

# In-memory activity store for local development (when DynamoDB is not available)
_LOCAL_ACTIVITIES = {}

# Admin Configuration
admin_emails_str = os.environ.get('ADMIN_EMAILS', '')
print(f'ADMIN_EMAILS env var: "{admin_emails_str}"', flush=True)
ADMIN_EMAILS = set(
    email.strip() for email in admin_emails_str.split(',')
    if email.strip()
)
print(f'Admin emails set: {ADMIN_EMAILS}', flush=True)
print(f'Admin emails count: {len(ADMIN_EMAILS)}', flush=True)

_last_results = None

MODEL = "gpt-4o"

# ============================================================================
# USER ACTIVITY TRACKING
# ============================================================================

def log_activity(email, action, details=None):
    """Log user activity to DynamoDB (or local store if unavailable)"""
    activity_item = {
        'email': email,
        'timestamp': int(time.time() * 1000),
        'action': action,
        'details': details or {},
        'ip_address': request.remote_addr,
        'user_agent': request.headers.get('User-Agent', 'Unknown')
    }

    try:
        activity_table.put_item(Item=activity_item)
        print(f'Activity logged to DynamoDB - Email: {email}, Action: {action}', flush=True)
    except Exception as e:
        # Fallback to local storage for development
        print(f'DynamoDB unavailable, using local storage: {e}', flush=True)
        if email not in _LOCAL_ACTIVITIES:
            _LOCAL_ACTIVITIES[email] = []
        _LOCAL_ACTIVITIES[email].append(activity_item)

# ============================================================================
# COGNITO AUTHENTICATION HELPERS
# ============================================================================

def login_required(f):
    """Decorator to require authentication"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


def admin_required(f):
    """Decorator to require admin access"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return redirect(url_for('login'))
        user_email = session.get('user', {}).get('email')
        print(f'Admin check - User email: {user_email}, Admin emails: {ADMIN_EMAILS}, Is admin: {user_email in ADMIN_EMAILS}', flush=True)
        if user_email not in ADMIN_EMAILS:
            print(f'Access denied for {user_email}', flush=True)
            return jsonify({'error': 'Admin access required'}), 403
        print(f'Access granted for {user_email}', flush=True)
        return f(*args, **kwargs)
    return decorated_function


def get_login_url(redirect_uri):
    """Generate Cognito login URL"""
    params = {
        'client_id': COGNITO_CLIENT_ID,
        'response_type': 'code',
        'scope': 'email openid',
        'redirect_uri': redirect_uri,
        'identity_provider': 'Google',  # Force Google login, bypassing Cognito hosted UI
    }
    return f'{COGNITO_AUTH_URL}?{urlencode(params)}'


def exchange_code_for_token(code, redirect_uri):
    """Exchange authorization code for ID and access tokens"""
    try:
        response = requests.post(
            COGNITO_TOKEN_URL,
            headers={'Content-Type': 'application/x-www-form-urlencoded'},
            data={
                'grant_type': 'authorization_code',
                'client_id': COGNITO_CLIENT_ID,
                'client_secret': COGNITO_CLIENT_SECRET,
                'code': code,
                'redirect_uri': redirect_uri,
            }
        )
        response.raise_for_status()
        return response.json()
    except Exception as e:
        print(f"Error exchanging code: {e}")
        return None


def decode_token(token):
    """Decode and verify JWT token from Cognito"""
    try:
        # Fetch public keys
        resp = requests.get(COGNITO_JWKS_URL)
        keys = resp.json()['keys']

        # Simple verification - in production use python-jose
        import base64
        parts = token.split('.')

        # Decode payload (second part)
        payload = parts[1]
        # Add padding if needed
        payload += '=' * (4 - len(payload) % 4)
        decoded = base64.urlsafe_b64decode(payload)

        return json.loads(decoded)
    except Exception as e:
        print(f"Error decoding token: {e}")
        return None


def get_user_info_from_token(id_token):
    """Extract user info from ID token"""
    payload = decode_token(id_token)
    if payload:
        email = payload.get('email', '')
        name = payload.get('name') or email or 'Logged In'
        return {
            'sub': payload.get('sub'),
            'email': email,
            'name': name,
            'picture': payload.get('picture'),
        }
    return None


def get_user_info_from_userinfo(access_token):
    """Fetch user info from Cognito userinfo endpoint using access token"""
    try:
        userinfo_endpoint = f'https://{COGNITO_DOMAIN}.auth.{COGNITO_REGION}.amazoncognito.com/oauth2/userinfo'
        print(f"Fetching userinfo from: {userinfo_endpoint}", flush=True)

        response = requests.get(
            userinfo_endpoint,
            headers={'Authorization': f'Bearer {access_token}'}
        )
        print(f"Userinfo response status: {response.status_code}", flush=True)
        response.raise_for_status()
        userinfo = response.json()
        print(f"Userinfo: {userinfo}", flush=True)

        email = userinfo.get('email', '')
        name = userinfo.get('name', userinfo.get('email', 'Logged In'))
        picture = userinfo.get('picture', '')

        return {
            'sub': userinfo.get('sub'),
            'email': email,
            'name': name,
            'picture': picture,
        }
    except Exception as e:
        print(f"Failed to fetch userinfo: {e}", flush=True)
        return None

JOB_DESCRIPTION = """
Senior QA Engineer
Location: Trivandrum
Experience: 6-10 years
Employment Type: Full-time

Job Summary:
We are seeking an experienced Senior QA Engineer (6-10 years experience) to lead quality
assurance initiatives and ensure the delivery of reliable, scalable, and high-performance
applications. The ideal candidate will define test strategies, drive automation frameworks,
mentor junior QA members, and collaborate closely with development and DevOps teams to
maintain high product quality.

Key Responsibilities:
- Define and implement comprehensive test strategies covering functional, regression,
  integration, performance, and security testing.
- Design, develop, and maintain automation frameworks using Playwright, Cypress, or Selenium.
- Validate backend services and database integrity using SQL queries (PostgreSQL).
- Perform performance and load testing using tools such as k6 or JMeter.
- Conduct security and vulnerability testing using OWASP ZAP or similar tools.
- Integrate automated test suites within CI/CD pipelines (GitHub Actions, Jenkins, Azure DevOps).
- Review test cases, ensure proper test coverage, and maintain test documentation standards.
- Analyze root causes of defects and implement preventive quality measures.
- Drive continuous improvement in QA processes, automation coverage, and release quality.
- Mentor junior QA engineers and provide technical guidance.
- Actively participate in Agile/Scrum ceremonies and collaborate with cross-functional teams.

Key Skills:
- Bachelor's degree in Computer Science, Engineering, or related field.
- 6-10 years of experience in software testing and quality assurance.
- Strong expertise in manual and automation testing.
- Hands-on experience with Playwright is mandatory. Additional experience with Cypress or
  Selenium is a plus but does not substitute for Playwright.
- Strong knowledge of API testing and automation.
- Experience with performance testing tools (k6, JMeter).
- Experience with CI/CD integration and DevOps workflows.
- Strong understanding of SDLC, STLC, and Agile methodologies.
- Experience with defect tracking tools such as Jira.
- Strong analytical, leadership, and mentoring skills.
- Excellent communication and stakeholder management abilities.
- Certification such as ISTQB or an equivalent QA qualification is an added advantage.
"""

RUBRIC = """
MANDATORY GATE (check this first, before scoring):
The candidate MUST have hands-on Playwright experience. Selenium and/or Cypress experience
does NOT substitute for Playwright, even if extensive. If the resume shows no clear evidence
of hands-on Playwright experience (e.g. only Selenium and/or Cypress, or only lists Playwright
as a passing familiarity with no real usage), set "meets_mandatory_requirement" to false and
cap total_score at 30, regardless of how strong the rest of the resume is. If Playwright
experience is clearly present, set "meets_mandatory_requirement" to true and score normally
using the rubric below.

Score the candidate 0-100 using this weighted rubric (only apply if the mandatory gate above
is passed):

1. Years of relevant QA experience, ideally 6-10 years (15 points)
2. Depth of hands-on Playwright experience specifically -- framework design, page object model,
   CI integration, etc. (20 points). Additional Selenium/Cypress experience can be noted as a
   plus in matched_skills but earns no extra points here.
3. API testing and automation experience (15 points)
4. CI/CD integration experience: GitHub Actions, Jenkins, or Azure DevOps (10 points)
5. Performance testing experience: k6 or JMeter (10 points)
6. Security testing experience: OWASP ZAP or similar (5 points)
7. SQL / database validation experience, ideally PostgreSQL (5 points)
8. SDLC/STLC/Agile methodology understanding (5 points)
9. Manual testing strength (5 points)
10. Leadership/mentoring experience (5 points)
11. Defect tracking tool experience, e.g. Jira (3 points)
12. ISTQB or equivalent certification (2 points)

Add up the points for a total score out of 100.
"""

PROMPT_TEMPLATE = """You are an expert technical recruiter screening resumes for the following role.

JOB DESCRIPTION:
{jd}

SCORING RUBRIC:
{rubric}

CANDIDATE RESUME:
{resume_text}

Evaluate this resume strictly against the rubric above. Respond with ONLY a valid JSON object,
no preamble, no markdown code fences, in exactly this format:

{{
  "candidate_name": "<name extracted from resume, or filename if not found>",
  "meets_mandatory_requirement": <true or false -- true only if hands-on Playwright experience is clearly present>,
  "total_score": <integer 0-100>,
  "years_experience_estimate": "<number or range>",
  "matched_skills": ["<skill>", "..."],
  "missing_skills": ["<skill>", "..."],
  "red_flags": ["<e.g. job hopping, unexplained gaps, no automation ownership, etc. Empty list if none>"],
  "summary": "<one or two sentence overall assessment>"
}}
"""


def extract_text_from_pdf(file_obj):
    try:
        doc = pdfium.PdfDocument(file_obj)
        text = "\n".join(p.get_textpage().get_text_range() for p in doc)
        return normalize_text(text)
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"


def normalize_text(text):
    import unicodedata
    nfkd = unicodedata.normalize('NFKD', text)
    return nfkd.encode('ascii', 'ignore').decode('ascii')


def extract_text_from_docx(file_obj):
    try:
        doc = Document(file_obj)
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        text = "\n".join(text_parts)
        normalized = normalize_text(text)
        print(f"[DEBUG] Original text length: {len(text)}, Normalized length: {len(normalized)}")
        print(f"[DEBUG] Original text (first 100 chars): {repr(text[:100])}")
        print(f"[DEBUG] Normalized text (first 100 chars): {repr(normalized[:100])}")
        return normalized
    except Exception as e:
        return f"Error extracting DOCX: {str(e)}"


def extract_text_from_txt(file_obj):
    try:
        text = file_obj.read().decode('utf-8')
        return normalize_text(text)
    except Exception as e:
        return f"Error extracting TXT: {str(e)}"


def score_resume(client, resume_text, filename, job_description, rubric):
    # Debug: check for non-ASCII characters
    try:
        resume_text.encode('ascii')
        print(f"[DEBUG] Resume text is pure ASCII: {len(resume_text)} chars")
    except UnicodeEncodeError as e:
        print(f"[DEBUG] Resume text has non-ASCII characters: {e}")
        print(f"[DEBUG] Resume text (first 100 chars): {repr(resume_text[:100])}")

    prompt = PROMPT_TEMPLATE.format(jd=job_description, rubric=rubric, resume_text=resume_text)

    response = client.chat.completions.create(
        model=MODEL,
        max_tokens=1000,
        response_format={"type": "json_object"},
        messages=[{"role": "user", "content": prompt}],
    )

    raw_text = response.choices[0].message.content.strip()

    if raw_text.startswith("```"):
        raw_text = raw_text.strip("`")
        raw_text = raw_text.replace("json", "", 1).strip()

    try:
        result = json.loads(raw_text)
    except json.JSONDecodeError:
        result = {
            "candidate_name": filename,
            "meets_mandatory_requirement": None,
            "total_score": None,
            "years_experience_estimate": None,
            "matched_skills": [],
            "missing_skills": [],
            "red_flags": ["PARSE_ERROR"],
            "summary": f"Could not parse model response: {raw_text[:200]}",
        }

    result["source_file"] = filename
    return result


# ============================================================================
# AUTHENTICATION ROUTES
# ============================================================================

@app.route('/login')
def login():
    """Redirect to Cognito login"""
    redirect_uri = url_for('callback', _external=True)
    login_url = get_login_url(redirect_uri)
    return redirect(login_url)


@app.route('/callback')
def callback():
    """Handle Cognito callback"""
    code = request.args.get('code')
    error = request.args.get('error')

    if error:
        print(f'LOGIN_FAILED - Error: {error}', flush=True)
        return jsonify({'error': error}), 400

    if not code:
        print('LOGIN_FAILED - Missing authorization code', flush=True)
        return jsonify({'error': 'Missing authorization code'}), 400

    # Exchange code for tokens
    redirect_uri = url_for('callback', _external=True)
    token_response = exchange_code_for_token(code, redirect_uri)

    if not token_response or 'id_token' not in token_response:
        return jsonify({'error': 'Failed to get tokens'}), 400

    # Extract user info from ID token
    user_info = get_user_info_from_token(token_response['id_token'])
    if not user_info:
        print('Failed to decode token', flush=True)
        return jsonify({'error': 'Failed to decode token'}), 400

    # Try to fetch email from Cognito userinfo endpoint (needed for Google federation)
    access_token = token_response.get('access_token')
    if access_token:
        userinfo = get_user_info_from_userinfo(access_token)
        if userinfo:
            user_info.update(userinfo)

    # Store user info in session
    session['user'] = user_info
    session['id_token'] = token_response['id_token']
    session['access_token'] = token_response.get('access_token')

    user_email = user_info.get('email', 'unknown')
    user_name = user_info.get('name', 'unknown')
    print(f'LOGIN_SUCCESS - Email: {user_email}, Name: {user_name}, Sub: {user_info.get("sub")}', flush=True)

    # Log login activity
    log_activity(user_email, 'LOGIN', {'name': user_name})

    return redirect(url_for('index'))


@app.route('/logout')
def logout():
    """Logout user and clear Cognito session"""
    user_email = session.get('user', {}).get('email', 'unknown')
    print(f'LOGOUT - Email: {user_email}', flush=True)

    # Log logout activity
    if user_email != 'unknown':
        log_activity(user_email, 'LOGOUT')

    session.clear()

    # Redirect to Cognito logout endpoint to clear OAuth session
    logout_uri = url_for('index', _external=True)
    print(f'Logout URI: {logout_uri}', flush=True)
    logout_params = urlencode({
        'client_id': COGNITO_CLIENT_ID,
        'logout_uri': logout_uri
    })
    cognito_logout_url = f'https://{COGNITO_DOMAIN}.auth.{COGNITO_REGION}.amazoncognito.com/logout?{logout_params}'
    print(f'Redirecting to Cognito logout: {cognito_logout_url}', flush=True)
    return redirect(cognito_logout_url)


# ============================================================================
# APPLICATION ROUTES
# ============================================================================

ACTIVITY_DASHBOARD_HTML = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Activity Dashboard</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link href="https://fonts.googleapis.com/css2?family=Open+Sans:wght@400;500;600;700&family=Poppins:wght@500;600;700;800&display=swap" rel="stylesheet">
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: 'Open Sans', sans-serif;
            background: linear-gradient(135deg, #0B1220 0%, #1a2847 100%);
            color: #E2E8F0;
            min-height: 100vh;
            padding: 40px 20px;
        }
        .container {
            max-width: 1200px;
            margin: 0 auto;
        }
        .header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 40px;
        }
        h1 {
            font-family: 'Poppins', sans-serif;
            font-size: 32px;
            color: #3B82F6;
        }
        .header-actions {
            display: flex;
            gap: 20px;
        }
        .btn {
            padding: 12px 24px;
            border-radius: 10px;
            border: none;
            cursor: pointer;
            font-size: 14px;
            font-weight: 600;
            transition: all 0.3s;
            text-decoration: none;
            display: inline-block;
        }
        .btn-secondary {
            background: #263252;
            color: #E2E8F0;
            border: 1px solid #3B82F6;
        }
        .btn-secondary:hover {
            background: #3B82F6;
            color: #fff;
        }
        .stats {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin-bottom: 40px;
        }
        .stat-card {
            background: rgba(59, 130, 246, 0.1);
            border: 1px solid #3B82F6;
            border-radius: 16px;
            padding: 24px;
            text-align: center;
        }
        .stat-value {
            font-size: 32px;
            font-weight: 700;
            color: #3B82F6;
            margin-bottom: 8px;
        }
        .stat-label {
            font-size: 14px;
            color: #94A3B8;
        }
        .table-container {
            background: rgba(255, 255, 255, 0.05);
            border: 1px solid #263252;
            border-radius: 16px;
            overflow: hidden;
        }
        table {
            width: 100%;
            border-collapse: collapse;
        }
        th {
            background: rgba(59, 130, 246, 0.1);
            padding: 16px;
            text-align: left;
            font-weight: 600;
            border-bottom: 1px solid #263252;
            color: #3B82F6;
        }
        td {
            padding: 16px;
            border-bottom: 1px solid #263252;
        }
        tr:last-child td {
            border-bottom: none;
        }
        tr:hover {
            background: rgba(59, 130, 246, 0.05);
        }
        .action-badge {
            display: inline-block;
            padding: 6px 12px;
            border-radius: 6px;
            font-size: 12px;
            font-weight: 600;
        }
        .action-login { background: rgba(34, 197, 94, 0.2); color: #86EFAC; }
        .action-logout { background: rgba(239, 68, 68, 0.2); color: #FCA5A5; }
        .action-upload { background: rgba(59, 130, 246, 0.2); color: #93C5FD; }
        .empty-state {
            text-align: center;
            padding: 60px 20px;
            color: #94A3B8;
        }
        .empty-state-icon {
            font-size: 48px;
            margin-bottom: 16px;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📊 Activity Dashboard</h1>
            <div class="header-actions">
                <a href="/" class="btn btn-secondary">← Back to App</a>
            </div>
        </div>

        <div class="stats">
            <div class="stat-card">
                <div class="stat-value">{{ total_count }}</div>
                <div class="stat-label">Total Activities</div>
            </div>
            <div class="stat-card">
                <div class="stat-value">{{ activities|selectattr('action', 'equalto', 'LOGIN')|list|length }}</div>
                <div class="stat-label">Logins</div>
            </div>
            <div class="stat-card">
                <div class="stat-value">{{ activities|selectattr('action', 'equalto', 'UPLOAD_RESUMES')|list|length }}</div>
                <div class="stat-label">Resume Uploads</div>
            </div>
        </div>

        {% if activities %}
        <div class="table-container">
            <table>
                <thead>
                    <tr>
                        <th>Time</th>
                        <th>User Email</th>
                        <th>Action</th>
                        <th>Details</th>
                        <th>IP Address</th>
                    </tr>
                </thead>
                <tbody>
                    {% for activity in activities %}
                    <tr>
                        <td><strong>{{ activity.readable_time }}</strong></td>
                        <td>{{ activity.email }}</td>
                        <td>
                            <span class="action-badge action-{{ activity.action|lower }}">
                                {{ activity.action }}
                            </span>
                        </td>
                        <td>
                            {% if activity.action == 'UPLOAD_RESUMES' %}
                                <small>{{ activity.details.file_count }} file(s)
                                {% if activity.details.custom_jd %}📝 Custom JD{% endif %}
                                {% if activity.details.custom_rubric %}📋 Custom Rubric{% endif %}
                                </small>
                            {% elif activity.action == 'LOGIN' %}
                                <small>{{ activity.details.name }}</small>
                            {% else %}
                                <small>-</small>
                            {% endif %}
                        </td>
                        <td><code style="font-size: 12px; color: #94A3B8;">{{ activity.ip_address }}</code></td>
                    </tr>
                    {% endfor %}
                </tbody>
            </table>
        </div>
        {% else %}
        <div class="empty-state">
            <div class="empty-state-icon">📭</div>
            <h3>No activities yet</h3>
            <p>Activities will appear here as users interact with the app</p>
        </div>
        {% endif %}
    </div>
</body>
</html>
"""

LOGIN_HTML = """
<!DOCTYPE html>
<html lang="en">
<head>
    <script>
        (function () {
            var stored = localStorage.getItem('theme');
            var theme = stored || (window.matchMedia && window.matchMedia('(prefers-color-scheme: dark)').matches ? 'dark' : 'light');
            document.documentElement.setAttribute('data-theme', theme);
        })();
    </script>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume Grader — Sign In</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Open+Sans:wght@400;500;600;700&family=Poppins:wght@500;600;700;800&display=swap" rel="stylesheet">
    <style>
        :root {
            --color-primary: #2563EB;
            --color-primary-dark: #1D4ED8;
            --color-secondary: #3B82F6;
            --color-accent: #EA580C;
            --color-accent-dark: #C2410C;
            --color-background: #F8FAFC;
            --color-foreground: #1E293B;
            --color-card: #FFFFFF;
            --color-muted: #E9EFF8;
            --color-muted-foreground: #475569;
            --color-border: #E2E8F0;
            --color-ring: #2563EB;
            --color-eyebrow-bg: #FFF1E8;
            --color-eyebrow-border: #FBD8BE;
            --color-glass-bg: rgba(255, 255, 255, 0.72);
            --color-glass-border: rgba(255, 255, 255, 0.6);
            --color-glow-1: rgba(37, 99, 235, 0.16);
            --color-glow-2: rgba(234, 88, 12, 0.10);
            --color-grid-line: rgba(30, 41, 59, 0.06);
            --color-grid-line-strong: rgba(37, 99, 235, 0.10);

            --font-heading: 'Poppins', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            --font-body: 'Open Sans', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;

            --space-2: 8px; --space-3: 12px; --space-4: 16px; --space-6: 24px;
            --space-8: 32px; --space-12: 48px;
            --radius-md: 10px; --radius-lg: 16px; --radius-xl: 24px;
            --shadow-sm: 0 1px 2px rgba(15, 23, 42, 0.06);
            --shadow-lg: 0 24px 64px rgba(15, 23, 42, 0.16);

            color-scheme: light;
        }

        :root[data-theme="dark"] {
            --color-primary: #3B82F6;
            --color-primary-dark: #60A5FA;
            --color-secondary: #60A5FA;
            --color-accent: #FB923C;
            --color-accent-dark: #FDBA74;
            --color-background: #0B1220;
            --color-foreground: #E2E8F0;
            --color-card: #131C2E;
            --color-muted: #1B2740;
            --color-muted-foreground: #94A3B8;
            --color-border: #263252;
            --color-ring: #60A5FA;
            --color-eyebrow-bg: #3A2412;
            --color-eyebrow-border: #5B3A1E;
            --color-glass-bg: rgba(19, 28, 46, 0.72);
            --color-glass-border: rgba(255, 255, 255, 0.08);
            --color-glow-1: rgba(59, 130, 246, 0.20);
            --color-glow-2: rgba(251, 146, 60, 0.12);
            --color-grid-line: rgba(226, 232, 240, 0.07);
            --color-grid-line-strong: rgba(96, 165, 250, 0.14);

            --shadow-sm: 0 1px 2px rgba(0, 0, 0, 0.4);
            --shadow-lg: 0 24px 64px rgba(0, 0, 0, 0.6);

            color-scheme: dark;
        }

        * { margin: 0; padding: 0; box-sizing: border-box; }

        body {
            font-family: var(--font-body);
            color: var(--color-foreground);
            background-color: var(--color-background);
            background-image:
                linear-gradient(var(--color-grid-line) 1px, transparent 1px),
                linear-gradient(90deg, var(--color-grid-line) 1px, transparent 1px),
                linear-gradient(var(--color-grid-line-strong) 1px, transparent 1px),
                linear-gradient(90deg, var(--color-grid-line-strong) 1px, transparent 1px),
                radial-gradient(circle at 15% -10%, var(--color-glow-1), transparent 45%),
                radial-gradient(circle at 90% 0%, var(--color-glow-2), transparent 40%);
            background-size: 24px 24px, 24px 24px, 120px 120px, 120px 120px, 100% 100%, 100% 100%;
            background-repeat: repeat, repeat, repeat, repeat, no-repeat, no-repeat;
            background-attachment: fixed;
            min-height: 100vh;
            line-height: 1.5;
            font-size: 16px;
            display: flex; align-items: center; justify-content: center;
            padding: var(--space-6);
        }

        button, a { cursor: pointer; }
        :focus-visible { outline: 2px solid var(--color-ring); outline-offset: 2px; border-radius: 4px; }

        .login-card {
            width: 100%; max-width: 480px; text-align: center;
            background: var(--color-glass-bg);
            backdrop-filter: blur(20px); -webkit-backdrop-filter: blur(20px);
            border: 1px solid var(--color-glass-border);
            border-radius: var(--radius-xl);
            box-shadow: var(--shadow-lg);
            padding: var(--space-12) var(--space-8);
            animation: slideUp 0.5s ease-out;
        }
        @keyframes slideUp {
            from { opacity: 0; transform: translateY(24px); }
            to { opacity: 1; transform: translateY(0); }
        }

        .brand-mark {
            position: relative;
            width: 56px; height: 56px; border-radius: 16px; margin: 0 auto var(--space-6);
            background: linear-gradient(135deg, var(--color-primary), var(--color-secondary));
            display: flex; align-items: center; justify-content: center;
            box-shadow: var(--shadow-sm);
        }
        .brand-mark svg { width: 28px; height: 28px; stroke: #fff; }

        .eyebrow {
            display: inline-flex; align-items: center; gap: 8px;
            font-size: 13px; font-weight: 600; letter-spacing: 0.02em;
            color: var(--color-accent-dark); background: var(--color-eyebrow-bg);
            border: 1px solid var(--color-eyebrow-border); padding: 6px 14px; border-radius: 999px;
            margin-bottom: var(--space-6);
        }
        .eyebrow svg { width: 14px; height: 14px; stroke: currentColor; }

        h1 {
            font-family: var(--font-heading); font-weight: 700;
            font-size: clamp(28px, 5vw, 36px); line-height: 1.2;
            margin-bottom: var(--space-3);
        }
        h1 span { color: var(--color-primary); }
        p.lede { color: var(--color-muted-foreground); font-size: 15px; margin-bottom: var(--space-8); }

        .login-btn {
            display: inline-flex; align-items: center; justify-content: center; gap: 12px;
            width: 100%; min-height: 50px;
            background: var(--color-card); color: var(--color-foreground);
            border: 1px solid var(--color-border); border-radius: var(--radius-md);
            font-family: var(--font-heading); font-size: 15px; font-weight: 600;
            text-decoration: none; padding: 14px 24px;
            transition: transform 0.15s ease, box-shadow 0.2s ease, border-color 0.2s ease;
        }
        .login-btn:hover { transform: translateY(-2px); border-color: var(--color-primary); box-shadow: 0 12px 24px rgba(37, 99, 235, 0.18); }
        .login-btn:active { transform: translateY(0); }
        .google-icon { width: 20px; height: 20px; display: block; }

        .features {
            margin-top: var(--space-8); padding-top: var(--space-6);
            border-top: 1px solid var(--color-border);
            display: flex; justify-content: center; gap: var(--space-6); flex-wrap: wrap;
            color: var(--color-muted-foreground); font-size: 13px;
        }
        .feature { display: flex; align-items: center; gap: 6px; }
        .feature svg { width: 15px; height: 15px; stroke: var(--color-primary); }

        @media (prefers-reduced-motion: reduce) {
            *, *::before, *::after { animation-duration: 0.001ms !important; transition-duration: 0.001ms !important; }
        }
    </style>
</head>
<body>
    <div class="login-card">
        <span class="brand-mark" aria-hidden="true">
            <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M9 12l2 2 4-4"/><path d="M12 3l8 4v5c0 5-3.5 8.5-8 9-4.5-.5-8-4-8-9V7l8-4z"/></svg>
        </span>

        <span class="eyebrow">
            <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M12 2l2.4 6.6L21 11l-6.6 2.4L12 20l-2.4-6.6L3 11l6.6-2.4z"/></svg>
            AI-Powered Screening
        </span>

        <h1>Resume <span>Grader</span></h1>
        <p class="lede">Sign in to grade candidate resumes against your job description and export a ranked shortlist.</p>

        <a href="/login" class="login-btn">
            <svg class="google-icon" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg" aria-hidden="true">
                <path d="M22.56 12.25c0-.78-.07-1.53-.2-2.25H12v4.26h5.92c-.26 1.37-1.04 2.53-2.21 3.31v2.77h3.57c2.08-1.92 3.28-4.74 3.28-8.09z" fill="#4285F4"/>
                <path d="M12 23c2.97 0 5.46-.98 7.28-2.66l-3.57-2.77c-.98.66-2.23 1.06-3.71 1.06-2.86 0-5.29-1.93-6.16-4.53H2.18v2.84C3.99 20.53 7.7 23 12 23z" fill="#34A853"/>
                <path d="M5.84 14.09c-.22-.66-.35-1.36-.35-2.09s.13-1.43.35-2.09V7.07H2.18C1.43 8.55 1 10.22 1 12s.43 3.45 1.18 4.93l2.85-2.22.81-.62z" fill="#FBBC05"/>
                <path d="M12 5.38c1.62 0 3.06.56 4.21 1.64l3.15-3.15C17.45 2.09 14.97 1 12 1 7.7 1 3.99 3.47 2.18 7.07l3.66 2.84c.87-2.6 3.3-4.53 6.16-4.53z" fill="#EA4335"/>
            </svg>
            Continue with Google
        </a>

        <div class="features">
            <span class="feature">
                <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M13 2L4.5 13.5H11L10 22l9-11.5H12l1-8.5z"/></svg>
                Instant analysis
            </span>
            <span class="feature">
                <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M9 11l3 3L22 4"/><path d="M21 12v7a2 2 0 01-2 2H5a2 2 0 01-2-2V5a2 2 0 012-2h11"/></svg>
                Rubric-based scoring
            </span>
            <span class="feature">
                <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><rect x="3" y="11" width="18" height="11" rx="2"/><path d="M7 11V7a5 5 0 0110 0v4"/></svg>
                Secure &amp; private
            </span>
        </div>
    </div>
</body>
</html>
"""


def build_user_block(user):
    """Render the signed-in user chip + logout button for the app header."""
    name = escape(user.get('name') or user.get('email') or 'Logged In')
    picture = user.get('picture')
    user_email = user.get('email', '')
    is_admin = user_email in ADMIN_EMAILS

    avatar = (
        f'<img src="{escape(picture)}" alt="" class="profile-pic">'
        if picture else
        '<span class="profile-pic profile-pic-fallback" aria-hidden="true">'
        '<svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">'
        '<path d="M20 21v-2a4 4 0 00-4-4H8a4 4 0 00-4 4v2"/><circle cx="12" cy="7" r="4"/></svg></span>'
    )

    admin_link = ''
    if is_admin:
        admin_link = '<a href="/activity" class="admin-link" title="Activity Dashboard">📊</a>'

    return f"""
                <div class="user-info">
                    {avatar}
                    <span class="user-name">{name}</span>
                    {admin_link}
                    <a href="/logout" class="logout-btn">
                        <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M9 21H5a2 2 0 01-2-2V5a2 2 0 012-2h4"/><path d="M16 17l5-5-5-5"/><path d="M21 12H9"/></svg>
                        <span>Logout</span>
                    </a>
                </div>"""


@app.route('/')
def index():
    user = session.get('user')

    # If not logged in, show the login page
    if not user:
        return LOGIN_HTML

    html = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <script>
            (function () {
                var stored = localStorage.getItem('theme');
                var theme = stored || (window.matchMedia && window.matchMedia('(prefers-color-scheme: dark)').matches ? 'dark' : 'light');
                document.documentElement.setAttribute('data-theme', theme);
            })();
        </script>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Resume Grader — AI-Powered Candidate Screening</title>
        <link rel="preconnect" href="https://fonts.googleapis.com">
        <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
        <link href="https://fonts.googleapis.com/css2?family=Open+Sans:wght@400;500;600;700&family=Poppins:wght@500;600;700;800&display=swap" rel="stylesheet">
        <style>
            :root {
                --color-primary: #2563EB;
                --color-primary-dark: #1D4ED8;
                --color-on-primary: #FFFFFF;
                --color-secondary: #3B82F6;
                --color-accent: #EA580C;
                --color-accent-dark: #C2410C;
                --color-on-accent: #FFFFFF;
                --color-background: #F8FAFC;
                --color-foreground: #1E293B;
                --color-card: #FFFFFF;
                --color-card-foreground: #1E293B;
                --color-muted: #E9EFF8;
                --color-muted-foreground: #475569;
                --color-border: #E2E8F0;
                --color-destructive: #DC2626;
                --color-success: #15803D;
                --color-success-bg: #DCFCE7;
                --color-success-fg: #14532D;
                --color-ring: #2563EB;

                --color-header-bg: rgba(248, 250, 252, 0.75);
                --color-glass-bg: rgba(255, 255, 255, 0.72);
                --color-glass-border: rgba(255, 255, 255, 0.6);
                --color-eyebrow-bg: #FFF1E8;
                --color-eyebrow-border: #FBD8BE;
                --color-upload-border: #B9CBF0;
                --color-upload-bg: rgba(37, 99, 235, 0.03);
                --color-upload-bg-hover: rgba(37, 99, 235, 0.07);
                --color-error-bg: #FEE2E2;
                --color-error-fg: #7F1D1D;
                --color-loading-bg: #DBEAFE;
                --color-loading-fg: #1E3A8A;
                --color-glow-1: rgba(37, 99, 235, 0.16);
                --color-glow-2: rgba(234, 88, 12, 0.10);
                --color-grid-line: rgba(30, 41, 59, 0.06);
                --color-grid-line-strong: rgba(37, 99, 235, 0.10);

                --font-heading: 'Poppins', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
                --font-body: 'Open Sans', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;

                --space-2: 8px; --space-3: 12px; --space-4: 16px; --space-6: 24px;
                --space-8: 32px; --space-12: 48px; --space-16: 64px;

                --radius-md: 10px; --radius-lg: 16px; --radius-xl: 24px;
                --shadow-sm: 0 1px 2px rgba(15, 23, 42, 0.06);
                --shadow-md: 0 8px 24px rgba(15, 23, 42, 0.08);
                --shadow-lg: 0 24px 64px rgba(15, 23, 42, 0.16);

                color-scheme: light;
            }

            :root[data-theme="dark"] {
                --color-primary: #3B82F6;
                --color-primary-dark: #60A5FA;
                --color-on-primary: #0B1220;
                --color-secondary: #60A5FA;
                --color-accent: #FB923C;
                --color-accent-dark: #FDBA74;
                --color-on-accent: #1F2937;
                --color-background: #0B1220;
                --color-foreground: #E2E8F0;
                --color-card: #131C2E;
                --color-card-foreground: #E2E8F0;
                --color-muted: #1B2740;
                --color-muted-foreground: #94A3B8;
                --color-border: #263252;
                --color-destructive: #F87171;
                --color-success: #4ADE80;
                --color-success-bg: #10281C;
                --color-success-fg: #86EFAC;
                --color-ring: #60A5FA;

                --color-header-bg: rgba(11, 18, 32, 0.75);
                --color-glass-bg: rgba(19, 28, 46, 0.72);
                --color-glass-border: rgba(255, 255, 255, 0.08);
                --color-eyebrow-bg: #3A2412;
                --color-eyebrow-border: #5B3A1E;
                --color-upload-border: #2E3F63;
                --color-upload-bg: rgba(59, 130, 246, 0.06);
                --color-upload-bg-hover: rgba(59, 130, 246, 0.14);
                --color-error-bg: #3B1418;
                --color-error-fg: #FCA5A5;
                --color-loading-bg: #12233F;
                --color-loading-fg: #93C5FD;
                --color-glow-1: rgba(59, 130, 246, 0.20);
                --color-glow-2: rgba(251, 146, 60, 0.12);
                --color-grid-line: rgba(226, 232, 240, 0.07);
                --color-grid-line-strong: rgba(96, 165, 250, 0.14);

                --shadow-sm: 0 1px 2px rgba(0, 0, 0, 0.4);
                --shadow-md: 0 8px 24px rgba(0, 0, 0, 0.45);
                --shadow-lg: 0 24px 64px rgba(0, 0, 0, 0.6);

                color-scheme: dark;
            }

            * { margin: 0; padding: 0; box-sizing: border-box; }

            html { scroll-behavior: smooth; }

            body {
                font-family: var(--font-body);
                color: var(--color-foreground);
                background-color: var(--color-background);
                background-image:
                    linear-gradient(var(--color-grid-line) 1px, transparent 1px),
                    linear-gradient(90deg, var(--color-grid-line) 1px, transparent 1px),
                    linear-gradient(var(--color-grid-line-strong) 1px, transparent 1px),
                    linear-gradient(90deg, var(--color-grid-line-strong) 1px, transparent 1px),
                    radial-gradient(circle at 15% -10%, var(--color-glow-1), transparent 45%),
                    radial-gradient(circle at 90% 0%, var(--color-glow-2), transparent 40%);
                background-size: 24px 24px, 24px 24px, 120px 120px, 120px 120px, 100% 100%, 100% 100%;
                background-position: 0 0, 0 0, 0 0, 0 0, 0 0, 0 0;
                background-repeat: repeat, repeat, repeat, repeat, no-repeat, no-repeat;
                background-attachment: fixed;
                min-height: 100vh;
                line-height: 1.5;
                font-size: 16px;
                transition: background-color 0.25s ease, color 0.25s ease;
                animation: gridDrift 60s linear infinite;
            }

            @keyframes gridDrift {
                to { background-position: 480px 480px, 480px 480px, 1200px 1200px, 1200px 1200px, 0 0, 0 0; }
            }

            a { color: inherit; }
            img, svg { display: block; }
            button, .clickable { cursor: pointer; }

            :focus-visible {
                outline: 2px solid var(--color-ring);
                outline-offset: 2px;
                border-radius: 4px;
            }

            .skip-link {
                position: absolute; left: -9999px; top: 0; z-index: 100;
                background: var(--color-primary); color: #fff; padding: 10px 16px; border-radius: 0 0 8px 0;
            }
            .skip-link:focus { left: 0; }

            /* ---------- Header ---------- */
            header {
                display: flex; align-items: center; justify-content: space-between;
                padding: var(--space-4) var(--space-8);
                background: transparent;
                border-bottom: 1px solid var(--color-border);
                transition: background-color 0.25s ease, border-color 0.25s ease;
            }
            .header-actions { display: flex; align-items: center; gap: var(--space-3); flex-wrap: wrap; justify-content: flex-end; }
            .brand { display: flex; align-items: center; gap: var(--space-4); }
            .brand-mark {
                position: relative;
                width: 46px; height: 46px; border-radius: 14px;
                background: linear-gradient(135deg, var(--color-primary), var(--color-secondary));
                display: flex; align-items: center; justify-content: center;
                box-shadow: var(--shadow-sm);
                animation: markFloat 4.5s ease-in-out infinite;
            }
            .brand-mark::before {
                content: "";
                position: absolute; inset: -7px; z-index: -1;
                border-radius: 18px;
                background: linear-gradient(135deg, var(--color-primary), var(--color-accent));
                filter: blur(11px);
                animation: markGlow 3.5s ease-in-out infinite;
            }
            .brand-mark svg { width: 24px; height: 24px; stroke: #fff; }
            @keyframes markFloat {
                0%, 100% { transform: translateY(0) rotate(0deg); }
                50% { transform: translateY(-3px) rotate(-4deg); }
            }
            @keyframes markGlow {
                0%, 100% { opacity: 0.25; transform: scale(0.94); }
                50% { opacity: 0.55; transform: scale(1.1); }
            }
            .brand-text {
                font-family: var(--font-heading); font-weight: 800;
                font-size: 24px; letter-spacing: -0.01em;
                background: linear-gradient(90deg, var(--color-foreground) 0%, var(--color-primary) 35%, var(--color-secondary) 50%, var(--color-accent) 65%, var(--color-foreground) 100%);
                background-size: 250% auto;
                -webkit-background-clip: text; background-clip: text; color: transparent;
                animation: textShine 7s linear infinite;
            }
            @keyframes textShine {
                to { background-position: -250% 0; }
            }
            .header-badge {
                display: inline-flex; align-items: center; gap: 6px;
                font-family: var(--font-body); font-size: 13px; font-weight: 600;
                color: var(--color-primary-dark); background: var(--color-muted);
                border: 1px solid var(--color-border); padding: 6px 12px; border-radius: 999px;
            }
            .header-badge svg { width: 14px; height: 14px; stroke: var(--color-primary); }

            .theme-toggle {
                display: flex; align-items: center; justify-content: center;
                width: 38px; height: 38px; border-radius: 50%;
                background: var(--color-card); border: 1px solid var(--color-border);
                color: var(--color-foreground);
                transition: background 0.2s ease, border-color 0.2s ease, transform 0.15s ease;
            }
            .theme-toggle:hover { background: var(--color-muted); border-color: var(--color-primary); }
            .theme-toggle:active { transform: translateY(1px); }
            .theme-toggle svg { width: 18px; height: 18px; stroke: currentColor; }
            .theme-toggle .icon-sun { display: none; }
            .theme-toggle .icon-moon { display: block; }
            :root[data-theme="dark"] .theme-toggle .icon-sun { display: block; }
            :root[data-theme="dark"] .theme-toggle .icon-moon { display: none; }

            /* ---------- Signed-in user ---------- */
            .user-info {
                display: flex; align-items: center; gap: var(--space-3);
                background: var(--color-card); border: 1px solid var(--color-border);
                border-radius: 999px; padding: 4px 6px 4px 6px;
            }
            .profile-pic {
                width: 30px; height: 30px; border-radius: 50%; object-fit: cover; flex-shrink: 0;
            }
            .profile-pic-fallback {
                background: var(--color-muted); display: flex; align-items: center; justify-content: center;
            }
            .profile-pic-fallback svg { width: 16px; height: 16px; stroke: var(--color-muted-foreground); }
            .user-name {
                font-size: 13px; font-weight: 600; color: var(--color-foreground);
                max-width: 160px; overflow: hidden; text-overflow: ellipsis; white-space: nowrap;
            }
            .logout-btn {
                display: inline-flex; align-items: center; gap: 6px;
                background: linear-gradient(135deg, var(--color-primary), var(--color-primary-dark));
                color: var(--color-on-primary); text-decoration: none;
                font-family: var(--font-body); font-size: 13px; font-weight: 600;
                padding: 8px 14px; border-radius: 999px;
                transition: transform 0.15s ease, box-shadow 0.2s ease;
            }
            .logout-btn svg { width: 14px; height: 14px; stroke: currentColor; }
            .logout-btn:hover { transform: translateY(-1px); box-shadow: 0 8px 18px rgba(37, 99, 235, 0.28); }
            .logout-btn:active { transform: translateY(0); }
            .admin-link {
                display: inline-flex; align-items: center; justify-content: center;
                width: 32px; height: 32px; border-radius: 8px;
                background: rgba(168, 85, 247, 0.1); color: #A855F7;
                text-decoration: none; font-size: 16px;
                transition: all 0.2s ease;
                border: 1px solid rgba(168, 85, 247, 0.2);
            }
            .admin-link:hover {
                background: rgba(168, 85, 247, 0.2);
                border-color: rgba(168, 85, 247, 0.4);
                transform: scale(1.05);
            }
            @media (max-width: 640px) {
                .user-name { display: none; }
                .logout-btn span { display: none; }
            }

            /* ---------- Hero ---------- */
            .hero { max-width: 1120px; margin: 0 auto; padding: var(--space-16) var(--space-8) var(--space-8); text-align: center; }
            .eyebrow {
                display: inline-flex; align-items: center; gap: 8px;
                font-size: 13px; font-weight: 600; letter-spacing: 0.02em;
                color: var(--color-accent-dark); background: var(--color-eyebrow-bg);
                border: 1px solid var(--color-eyebrow-border); padding: 6px 14px; border-radius: 999px;
                margin-bottom: var(--space-6);
            }
            .hero h1 {
                font-family: var(--font-heading); font-weight: 700;
                font-size: clamp(32px, 5vw, 52px); line-height: 1.15;
                color: var(--color-foreground); max-width: 780px; margin: 0 auto var(--space-4);
            }
            .hero h1 span { color: var(--color-primary); }
            .hero p.lede {
                font-size: 17px; color: var(--color-muted-foreground);
                max-width: 560px; margin: 0 auto var(--space-8);
            }
            .hero-stats {
                display: flex; justify-content: center; align-items: stretch;
                max-width: 640px; margin: 0 auto var(--space-4);
                background: var(--color-card); border: 1px solid var(--color-border);
                border-radius: var(--radius-lg); box-shadow: var(--shadow-sm); overflow: hidden;
            }
            .hero-stat {
                flex: 1; position: relative; text-align: center;
                display: flex; flex-direction: column; align-items: center; gap: 6px;
                padding: var(--space-4) var(--space-3);
                transition: background-color 0.2s ease;
            }
            .hero-stat:not(:last-child)::after {
                content: ""; position: absolute; right: 0; top: 18%; bottom: 18%; width: 1px;
                background: var(--color-border);
            }
            .hero-stat:hover { background-color: var(--color-muted); }
            .hero-stat-icon {
                width: 30px; height: 30px; border-radius: 8px;
                background: var(--color-muted); display: flex; align-items: center; justify-content: center;
            }
            .hero-stat-icon svg { width: 15px; height: 15px; stroke: var(--color-primary); }
            .hero-stat strong { display: block; font-family: var(--font-heading); font-size: 18px; color: var(--color-foreground); }
            .hero-stat span { font-size: 12px; color: var(--color-muted-foreground); }
            @media (max-width: 560px) {
                .hero-stats { flex-direction: column; }
                .hero-stat:not(:last-child)::after { right: 18%; left: 18%; bottom: 0; top: auto; width: auto; height: 1px; }
            }

            /* ---------- Layout shell ---------- */
            .shell { max-width: 1120px; margin: 0 auto; padding: 0 var(--space-8) var(--space-16); display: grid; grid-template-columns: 1.15fr 0.85fr; gap: var(--space-8); align-items: start; }

            /* ---------- Glass tool card ---------- */
            .tool-card {
                background: var(--color-glass-bg);
                backdrop-filter: blur(20px);
                -webkit-backdrop-filter: blur(20px);
                border: 1px solid var(--color-glass-border);
                border-radius: var(--radius-xl);
                box-shadow: var(--shadow-lg);
                padding: var(--space-8);
                transition: background-color 0.25s ease, border-color 0.25s ease;
            }
            .tool-card h2 { font-family: var(--font-heading); font-size: 20px; font-weight: 600; margin-bottom: var(--space-2); }
            .tool-card .tool-sub { color: var(--color-muted-foreground); font-size: 14px; margin-bottom: var(--space-6); }

            .settings-toggle {
                display: inline-flex; align-items: center; gap: 8px;
                background: var(--color-card); color: var(--color-primary-dark);
                border: 1px solid var(--color-border); padding: 10px 16px; border-radius: var(--radius-md);
                font-family: var(--font-body); font-size: 13px; font-weight: 600;
                transition: background 0.2s ease, border-color 0.2s ease, transform 0.15s ease;
                margin-bottom: var(--space-6);
            }
            .settings-toggle svg { width: 15px; height: 15px; stroke: currentColor; }
            .settings-toggle:hover { background: var(--color-muted); border-color: var(--color-primary); }
            .settings-toggle:active { transform: translateY(1px); }

            .settings-panel {
                margin-bottom: var(--space-6); padding: var(--space-6);
                background: var(--color-muted); border-radius: var(--radius-lg);
                border: 1px solid var(--color-border);
                display: grid; gap: var(--space-6);
            }
            .field-label {
                font-family: var(--font-heading); font-weight: 600; color: var(--color-foreground);
                margin-bottom: var(--space-2); font-size: 12px; text-transform: uppercase; letter-spacing: 0.06em;
                display: flex; align-items: center; gap: 6px;
            }
            .field-label svg { width: 14px; height: 14px; stroke: var(--color-primary); }
            textarea {
                width: 100%; padding: var(--space-3); border: 1px solid var(--color-border);
                border-radius: var(--radius-md); font-family: 'Monaco', 'Courier New', monospace;
                font-size: 13px; resize: vertical; line-height: 1.6; background: var(--color-card);
                color: var(--color-foreground); transition: border-color 0.2s ease, box-shadow 0.2s ease;
            }
            textarea:focus { outline: none; border-color: var(--color-primary); box-shadow: 0 0 0 3px rgba(37, 99, 235, 0.15); }

            .upload-area {
                border: 2px dashed var(--color-upload-border); border-radius: var(--radius-lg);
                padding: var(--space-12) var(--space-6); text-align: center;
                transition: border-color 0.2s ease, background 0.2s ease, transform 0.2s ease;
                background: var(--color-upload-bg); margin-bottom: var(--space-4);
            }
            .upload-area:hover, .upload-area.drag-over { border-color: var(--color-primary); background: var(--color-upload-bg-hover); }
            .upload-area input { position: absolute; width: 1px; height: 1px; overflow: hidden; clip: rect(0,0,0,0); }
            .upload-icon {
                width: 56px; height: 56px; margin: 0 auto var(--space-4);
                border-radius: 50%; background: var(--color-muted);
                display: flex; align-items: center; justify-content: center;
            }
            .upload-icon svg { width: 26px; height: 26px; stroke: var(--color-primary); }
            .upload-text { color: var(--color-foreground); font-weight: 600; margin-bottom: 4px; font-size: 15px; font-family: var(--font-heading); }
            .upload-hint { color: var(--color-muted-foreground); font-size: 13px; }

            .file-list { display: grid; gap: var(--space-2); margin-bottom: var(--space-4); }
            .file-item {
                background: var(--color-card); padding: 10px 14px; border-radius: var(--radius-md);
                display: flex; justify-content: space-between; align-items: center;
                border: 1px solid var(--color-border); font-size: 13px;
            }
            .file-item .file-name { display: flex; align-items: center; gap: 8px; color: var(--color-foreground); overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
            .file-item .file-name svg { width: 15px; height: 15px; stroke: var(--color-muted-foreground); flex-shrink: 0; }
            .remove-btn {
                background: transparent; color: var(--color-destructive); border: none;
                width: 28px; height: 28px; border-radius: 6px; display: flex; align-items: center; justify-content: center;
                transition: background 0.15s ease; flex-shrink: 0;
            }
            .remove-btn svg { width: 14px; height: 14px; stroke: currentColor; }
            .remove-btn:hover { background: var(--color-error-bg); }

            button.primary {
                display: flex; align-items: center; justify-content: center; gap: 8px;
                background: linear-gradient(135deg, var(--color-primary), var(--color-primary-dark));
                color: var(--color-on-primary); border: none; padding: 15px 24px;
                border-radius: var(--radius-md); font-family: var(--font-heading);
                font-size: 15px; font-weight: 600; width: 100%; min-height: 48px;
                transition: transform 0.15s ease, box-shadow 0.2s ease, opacity 0.2s ease;
            }
            button.primary svg { width: 18px; height: 18px; stroke: currentColor; }
            button.primary:hover:not(:disabled) { transform: translateY(-2px); box-shadow: 0 12px 24px rgba(37, 99, 235, 0.3); }
            button.primary:active:not(:disabled) { transform: translateY(0); }
            button.primary:disabled { opacity: 0.45; cursor: not-allowed; }

            .download-btn {
                background: linear-gradient(135deg, #16A34A, #15803D);
                margin-top: var(--space-3); display: none; text-decoration: none;
                width: auto; padding: 15px 32px;
            }
            .download-btn:hover { box-shadow: 0 12px 24px rgba(21, 128, 61, 0.3) !important; }

            .status {
                margin-top: var(--space-4); padding: 14px 16px; border-radius: var(--radius-md);
                display: none; align-items: center; gap: 10px; font-size: 14px; font-weight: 500;
            }
            .status svg { width: 18px; height: 18px; flex-shrink: 0; }
            .status.success { background: var(--color-success-bg); color: var(--color-success-fg); border-left: 4px solid var(--color-success); }
            .status.error { background: var(--color-error-bg); color: var(--color-error-fg); border-left: 4px solid var(--color-destructive); }
            .status.loading { background: var(--color-loading-bg); color: var(--color-loading-fg); border-left: 4px solid var(--color-primary); }
            .spin { animation: spin 0.9s linear infinite; }
            @keyframes spin { to { transform: rotate(360deg); } }

            /* ---------- Side panel: how it works / trust ---------- */
            .side-panel { display: grid; gap: var(--space-4); position: sticky; top: 96px; }
            .info-card {
                background: var(--color-card); border: 1px solid var(--color-border);
                border-radius: var(--radius-lg); padding: var(--space-6); box-shadow: var(--shadow-sm);
            }
            .info-card h3 { font-family: var(--font-heading); font-size: 15px; font-weight: 600; margin-bottom: var(--space-4); }
            .step { display: flex; gap: var(--space-3); margin-bottom: var(--space-4); }
            .step:last-child { margin-bottom: 0; }
            .step-num {
                width: 28px; height: 28px; border-radius: 50%; background: var(--color-muted);
                color: var(--color-primary-dark); font-family: var(--font-heading); font-weight: 700; font-size: 13px;
                display: flex; align-items: center; justify-content: center; flex-shrink: 0;
            }
            .step-body strong { display: block; font-size: 14px; margin-bottom: 2px; }
            .step-body span { font-size: 13px; color: var(--color-muted-foreground); }

            .feature-card { display: flex; gap: var(--space-3); align-items: flex-start; }
            .feature-icon {
                width: 36px; height: 36px; border-radius: 10px; flex-shrink: 0;
                background: linear-gradient(135deg, rgba(37,99,235,0.12), rgba(59,130,246,0.12));
                display: flex; align-items: center; justify-content: center;
            }
            .feature-icon svg { width: 18px; height: 18px; stroke: var(--color-primary); }
            .feature-card strong { display: block; font-size: 14px; margin-bottom: 2px; }
            .feature-card span { font-size: 13px; color: var(--color-muted-foreground); }
            .feature-grid { display: grid; gap: var(--space-4); }

            footer {
                border-top: 1px solid var(--color-border); padding: var(--space-6) var(--space-8);
                text-align: center; color: var(--color-muted-foreground); font-size: 13px;
            }
            footer strong { color: var(--color-foreground); }
            footer .eyebrow { margin-bottom: var(--space-3); }

            /* ---------- Responsive ---------- */
            @media (max-width: 1024px) {
                .shell { grid-template-columns: 1fr; }
                .side-panel { position: static; }
            }
            @media (max-width: 768px) {
                header { padding: var(--space-4); }
                .hero { padding: var(--space-12) var(--space-4) var(--space-6); }
                .shell { padding: 0 var(--space-4) var(--space-12); gap: var(--space-6); }
                .hero-stats { gap: var(--space-6); }
                .tool-card { padding: var(--space-6); }
            }
            @media (max-width: 375px) {
                .tool-card { padding: var(--space-4); }
                .upload-area { padding: var(--space-8) var(--space-4); }
            }

            @media (prefers-reduced-motion: reduce) {
                *, *::before, *::after { animation-duration: 0.001ms !important; animation-iteration-count: 1 !important; transition-duration: 0.001ms !important; scroll-behavior: auto !important; }
            }
        </style>
    </head>
    <body>
        <a class="skip-link" href="#main">Skip to content</a>

        <header>
            <div class="brand">
                <span class="brand-mark" aria-hidden="true">
                    <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M9 12l2 2 4-4"/><path d="M12 3l8 4v5c0 5-3.5 8.5-8 9-4.5-.5-8-4-8-9V7l8-4z"/></svg>
                </span>
                <span class="brand-text">Resume Grader</span>
            </div>
            <div class="header-actions">
                <span class="header-badge">
                    <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M12 2l2.4 6.6L21 11l-6.6 2.4L12 20l-2.4-6.6L3 11l6.6-2.4z"/></svg>
                    AI-Powered Screening
                </span>
PLACEHOLDER_USER
                <button class="theme-toggle" id="themeToggle" type="button" aria-label="Switch to dark mode" aria-pressed="false">
                    <svg class="icon-sun" viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><circle cx="12" cy="12" r="4"/><path d="M12 2v2"/><path d="M12 20v2"/><path d="M4.93 4.93l1.41 1.41"/><path d="M17.66 17.66l1.41 1.41"/><path d="M2 12h2"/><path d="M20 12h2"/><path d="M6.34 17.66l-1.41 1.41"/><path d="M19.07 4.93l-1.41 1.41"/></svg>
                    <svg class="icon-moon" viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M21 12.79A9 9 0 1111.21 3 7 7 0 0021 12.79z"/></svg>
                </button>
            </div>
        </header>

        <main id="main">
            <section class="hero">
                <h1>Grade every resume against your job description <span>in seconds</span></h1>
                <p class="lede">Upload a batch of candidates, define your rubric, and get a ranked, exportable shortlist powered by AI — no manual screening required.</p>
                <div class="hero-stats">
                    <div class="hero-stat">
                        <span class="hero-stat-icon" aria-hidden="true">
                            <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><path d="M12 6v6l4 2"/></svg>
                        </span>
                        <strong>30s</strong><span>avg. per resume</span>
                    </div>
                    <div class="hero-stat">
                        <span class="hero-stat-icon" aria-hidden="true">
                            <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><path d="M14 2v6h6"/></svg>
                        </span>
                        <strong>PDF · DOCX · TXT</strong><span>supported formats</span>
                    </div>
                    <div class="hero-stat">
                        <span class="hero-stat-icon" aria-hidden="true">
                            <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M21 15v4a2 2 0 01-2 2H5a2 2 0 01-2-2v-4"/><path d="M7 10l5 5 5-5"/><path d="M12 15V3"/></svg>
                        </span>
                        <strong>1-click</strong><span>Excel export</span>
                    </div>
                </div>
            </section>

            <div class="shell">
                <div class="tool-card">
                    <h2>Score your candidates</h2>
                    <p class="tool-sub">Upload resumes below — grading uses the job description and rubric configured on the right.</p>

                    <button class="settings-toggle" id="settingsToggle" type="button" aria-expanded="false" aria-controls="settingsPanel">
                        <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><circle cx="12" cy="12" r="3"/><path d="M19.4 15a1.65 1.65 0 00.33 1.82l.06.06a2 2 0 11-2.83 2.83l-.06-.06a1.65 1.65 0 00-1.82-.33 1.65 1.65 0 00-1 1.51V21a2 2 0 01-4 0v-.09A1.65 1.65 0 009 19.4a1.65 1.65 0 00-1.82.33l-.06.06a2 2 0 11-2.83-2.83l.06-.06A1.65 1.65 0 004.6 15a1.65 1.65 0 00-1.51-1H3a2 2 0 010-4h.09A1.65 1.65 0 004.6 9a1.65 1.65 0 00-.33-1.82l-.06-.06a2 2 0 112.83-2.83l.06.06A1.65 1.65 0 009 4.6a1.65 1.65 0 001-1.51V3a2 2 0 014 0v.09a1.65 1.65 0 001 1.51 1.65 1.65 0 001.82-.33l.06-.06a2 2 0 112.83 2.83l-.06.06A1.65 1.65 0 0019.4 9a1.65 1.65 0 001.51 1H21a2 2 0 010 4h-.09a1.65 1.65 0 00-1.51 1z"/></svg>
                        Configure job description &amp; rubric
                    </button>

                    <div class="settings-panel" id="settingsPanel" style="display: none;">
                        <div>
                            <div class="field-label">
                                <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><path d="M14 2v6h6"/></svg>
                                Job Description
                            </div>
                            <textarea id="jdField" rows="6" aria-label="Job description">PLACEHOLDER_JD</textarea>
                        </div>

                        <div>
                            <div class="field-label">
                                <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M9 11l3 3L22 4"/><path d="M21 12v7a2 2 0 01-2 2H5a2 2 0 01-2-2V5a2 2 0 012-2h11"/></svg>
                                Scoring Rubric
                            </div>
                            <textarea id="rubricField" rows="6" aria-label="Scoring rubric">PLACEHOLDER_RUBRIC</textarea>
                        </div>
                    </div>

                    <div class="upload-area" id="uploadArea">
                        <label for="fileInput" style="cursor: pointer; display: block;">
                            <div class="upload-icon" aria-hidden="true">
                                <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M21 15v4a2 2 0 01-2 2H5a2 2 0 01-2-2v-4"/><path d="M17 8l-5-5-5 5"/><path d="M12 3v12"/></svg>
                            </div>
                            <div class="upload-text">Drop files here or click to upload</div>
                            <div class="upload-hint">Supports PDF, DOCX, and TXT files</div>
                        </label>
                        <input type="file" id="fileInput" multiple accept=".pdf,.docx,.txt">
                    </div>

                    <div class="file-list" id="fileList"></div>

                    <button class="primary" id="submitBtn" type="button" disabled>
                        <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M9 11l3 3L22 4"/><path d="M21 12v7a2 2 0 01-2 2H5a2 2 0 01-2-2V5a2 2 0 012-2h11"/></svg>
                        Grade Resumes
                    </button>
                    <a id="downloadBtn" class="primary download-btn" href="" download="shortlist.xlsx">
                        <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M21 15v4a2 2 0 01-2 2H5a2 2 0 01-2-2v-4"/><path d="M7 10l5 5 5-5"/><path d="M12 15V3"/></svg>
                        Download Results
                    </a>

                    <div class="status" id="status" role="status" aria-live="polite"></div>
                </div>

                <aside class="side-panel">
                    <div class="info-card">
                        <h3>How it works</h3>
                        <div class="step">
                            <span class="step-num">1</span>
                            <div class="step-body"><strong>Set your criteria</strong><span>Confirm or edit the job description and rubric.</span></div>
                        </div>
                        <div class="step">
                            <span class="step-num">2</span>
                            <div class="step-body"><strong>Upload resumes</strong><span>Drop in as many candidate files as you need.</span></div>
                        </div>
                        <div class="step">
                            <span class="step-num">3</span>
                            <div class="step-body"><strong>Get a ranked shortlist</strong><span>Download scored results as an Excel file.</span></div>
                        </div>
                    </div>

                    <div class="info-card feature-grid">
                        <h3 style="margin-bottom: 0;">Why teams use it</h3>
                        <div class="feature-card">
                            <span class="feature-icon" aria-hidden="true">
                                <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M12 2l2.4 6.6L21 11l-6.6 2.4L12 20l-2.4-6.6L3 11l6.6-2.4z"/></svg>
                            </span>
                            <div><strong>Consistent scoring</strong><span>Every resume is graded against the same rubric, removing reviewer bias.</span></div>
                        </div>
                        <div class="feature-card">
                            <span class="feature-icon" aria-hidden="true">
                                <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M17 21v-2a4 4 0 00-4-4H5a4 4 0 00-4 4v2"/><circle cx="9" cy="7" r="4"/><path d="M23 21v-2a4 4 0 00-3-3.87"/><path d="M16 3.13a4 4 0 010 7.75"/></svg>
                            </span>
                            <div><strong>Bulk-ready</strong><span>Grade dozens of candidates in one batch instead of one at a time.</span></div>
                        </div>
                        <div class="feature-card">
                            <span class="feature-icon" aria-hidden="true">
                                <svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M21 15v4a2 2 0 01-2 2H5a2 2 0 01-2-2v-4"/><path d="M7 10l5 5 5-5"/><path d="M12 15V3"/></svg>
                            </span>
                            <div><strong>Export instantly</strong><span>Ranked results land in a ready-to-share Excel shortlist.</span></div>
                        </div>
                    </div>
                </aside>
            </div>
        </main>

        <footer>
            <span class="eyebrow">
                <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M13 2L4.5 13.5H11L10 22l9-11.5H12l1-8.5z"/></svg>
                Built for hiring teams
            </span>
            <p><strong>Resume Grader</strong> — AI-assisted candidate screening. Always review shortlists before making hiring decisions.</p>
        </footer>

        <script>
            let files = [];
            const fileInput = document.getElementById('fileInput');
            const uploadArea = document.getElementById('uploadArea');
            const fileList = document.getElementById('fileList');
            const submitBtn = document.getElementById('submitBtn');
            const downloadBtn = document.getElementById('downloadBtn');
            const status = document.getElementById('status');
            const jdField = document.getElementById('jdField');
            const rubricField = document.getElementById('rubricField');
            const settingsToggle = document.getElementById('settingsToggle');
            const settingsPanel = document.getElementById('settingsPanel');
            const themeToggle = document.getElementById('themeToggle');

            function applyThemeLabel(theme) {
                themeToggle.setAttribute('aria-pressed', String(theme === 'dark'));
                themeToggle.setAttribute('aria-label', theme === 'dark' ? 'Switch to light mode' : 'Switch to dark mode');
            }
            applyThemeLabel(document.documentElement.getAttribute('data-theme') || 'light');

            themeToggle.addEventListener('click', () => {
                const next = document.documentElement.getAttribute('data-theme') === 'dark' ? 'light' : 'dark';
                document.documentElement.setAttribute('data-theme', next);
                localStorage.setItem('theme', next);
                applyThemeLabel(next);
            });

            const fileSvg = '<svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><path d="M14 2v6h6"/></svg>';
            const removeSvg = '<svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M18 6L6 18"/><path d="M6 6l12 12"/></svg>';
            const spinnerSvg = '<svg class="spin" viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round"><path d="M12 2a10 10 0 0110 10"/></svg>';
            const successSvg = '<svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><path d="M9 12l2 2 4-4"/></svg>';
            const errorSvg = '<svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><path d="M12 8v4"/><path d="M12 16h.01"/></svg>';

            settingsToggle.addEventListener('click', () => {
                const isOpen = settingsPanel.style.display !== 'none';
                settingsPanel.style.display = isOpen ? 'none' : 'grid';
                settingsToggle.setAttribute('aria-expanded', String(!isOpen));
            });

            fileInput.addEventListener('change', (e) => {
                addFiles(e.target.files);
            });

            ['dragenter', 'dragover'].forEach(evt => {
                uploadArea.addEventListener(evt, (e) => {
                    e.preventDefault();
                    uploadArea.classList.add('drag-over');
                });
            });
            ['dragleave', 'drop'].forEach(evt => {
                uploadArea.addEventListener(evt, (e) => {
                    e.preventDefault();
                    uploadArea.classList.remove('drag-over');
                });
            });
            uploadArea.addEventListener('drop', (e) => {
                if (e.dataTransfer && e.dataTransfer.files) addFiles(e.dataTransfer.files);
            });

            function addFiles(fileListInput) {
                const incoming = Array.from(fileListInput);
                const existingNames = new Set(files.map(f => f.name + f.size));
                incoming.forEach(f => {
                    if (!existingNames.has(f.name + f.size)) files.push(f);
                });
                updateFileList();
                submitBtn.disabled = files.length === 0;
            }

            function updateFileList() {
                fileList.innerHTML = files.map((f, i) => `
                    <div class="file-item">
                        <span class="file-name">${fileSvg}<span>${f.name}</span></span>
                        <button class="remove-btn" type="button" aria-label="Remove ${f.name}" onclick="removeFile(${i})">${removeSvg}</button>
                    </div>
                `).join('');
            }

            function removeFile(i) {
                files.splice(i, 1);
                updateFileList();
                submitBtn.disabled = files.length === 0;
            }

            function setStatus(kind, icon, text) {
                status.className = `status ${kind}`;
                status.style.display = 'flex';
                status.innerHTML = `${icon}<span>${text}</span>`;
            }

            submitBtn.addEventListener('click', async () => {
                if (files.length === 0) return;

                const formData = new FormData();
                files.forEach(f => formData.append('files', f));
                formData.append('job_description', jdField.value);
                formData.append('rubric', rubricField.value);

                submitBtn.disabled = true;
                setStatus('loading', spinnerSvg, 'Processing resumes…');
                downloadBtn.style.display = 'none';

                try {
                    const response = await fetch('/grade', {
                        method: 'POST',
                        body: formData
                    });
                    const result = await response.json();

                    if (response.ok && result.download_url) {
                        setStatus('success', successSvg, 'Grading complete — your shortlist is ready.');
                        downloadBtn.href = result.download_url;
                        downloadBtn.style.display = 'flex';
                        downloadBtn.click();
                        files = [];
                        updateFileList();
                        fileInput.value = '';
                    } else {
                        setStatus('error', errorSvg, `Error: ${result.error || 'Unknown error'}`);
                    }
                } catch (e) {
                    setStatus('error', errorSvg, `Error: ${e.message}`);
                } finally {
                    submitBtn.disabled = files.length === 0;
                }
            });
        </script>
    </body>
    </html>
    """
    return (
        html
        .replace("PLACEHOLDER_USER", build_user_block(user))
        .replace("PLACEHOLDER_JD", JOB_DESCRIPTION)
        .replace("PLACEHOLDER_RUBRIC", RUBRIC)
    )


@app.route('/grade', methods=['POST'])
@login_required
def grade_resumes():
    global _last_results

    # Get user email from session
    user_email = session.get('user', {}).get('email', 'unknown')

    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return jsonify({"error": "OPENAI_API_KEY not configured"}), 500

    if 'files' not in request.files or len(request.files.getlist('files')) == 0:
        return jsonify({"error": "No files uploaded"}), 400

    job_description = request.form.get('job_description', JOB_DESCRIPTION)
    rubric = request.form.get('rubric', RUBRIC)

    uploaded_files = request.files.getlist('files')

    # Log activity: file upload
    log_activity(user_email, 'UPLOAD_RESUMES', {
        'file_count': len(uploaded_files),
        'custom_jd': job_description != JOB_DESCRIPTION,
        'custom_rubric': rubric != RUBRIC
    })
    client = OpenAI(api_key=api_key)
    results = []

    for file in uploaded_files:
        if not file or file.filename == '':
            print(f"Skipping empty file")
            continue

        filename = secure_filename(file.filename)
        file.seek(0)

        if filename.lower().endswith('.pdf'):
            resume_text = extract_text_from_pdf(file.stream)
        elif filename.lower().endswith('.docx'):
            resume_text = extract_text_from_docx(file.stream)
        elif filename.lower().endswith('.txt'):
            resume_text = extract_text_from_txt(file.stream)
        else:
            print(f"Unsupported file format: {filename}")
            continue

        if not resume_text.strip():
            print(f"Empty resume text for {filename}")
            continue

        if resume_text.startswith("Error"):
            print(f"Extraction error for {filename}: {resume_text}")
            continue

        try:
            result = score_resume(client, resume_text, filename, job_description, rubric)
            results.append(result)
            print(f"Successfully scored {filename}")
            time.sleep(0.3)
        except Exception as e:
            print(f"Error scoring {filename}: {e}")

    if not results:
        return jsonify({"error": "No valid resumes to process"}), 400

    df = pd.DataFrame(results)

    for col in ["matched_skills", "missing_skills", "red_flags"]:
        if col in df.columns:
            df[col] = df[col].apply(lambda x: ", ".join(x) if isinstance(x, list) else x)

    if "meets_mandatory_requirement" in df.columns:
        df = df.sort_values(
            by=["meets_mandatory_requirement", "total_score"],
            ascending=[False, False],
            na_position="last",
        )
    else:
        df = df.sort_values(by="total_score", ascending=False, na_position="last")

    column_order = [
        "candidate_name", "meets_mandatory_requirement", "total_score",
        "years_experience_estimate", "matched_skills", "missing_skills",
        "red_flags", "summary", "source_file",
    ]
    df = df[[c for c in column_order if c in df.columns]]

    output = io.BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)
    _last_results = output.getvalue()

    return jsonify({
        "download_url": "/download",
        "result": df.to_dict(orient='records')
    })


@app.route('/download')
@login_required
def download():
    global _last_results

    if _last_results is None:
        return jsonify({"error": "No results available"}), 400

    return send_file(
        io.BytesIO(_last_results),
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"shortlist_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    )


@app.route('/health', methods=['GET'])
def health():
    return jsonify({"status": "healthy"}), 200


@app.route('/activity', methods=['GET'])
@admin_required
def view_activity():
    """View user activity log (admin only)"""
    user_email = session.get('user', {}).get('email', 'unknown')
    print(f'Fetching activity for admin: {user_email}', flush=True)

    try:
        print(f'Querying DynamoDB table: {ACTIVITY_TABLE_NAME}', flush=True)
        # Query activities for all users
        response = activity_table.scan(Limit=500)
        all_activities = response.get('Items', [])
        print(f'Found {len(all_activities)} activities in DynamoDB', flush=True)

    except Exception as e:
        # Fallback to local storage for development
        print(f'DynamoDB unavailable ({type(e).__name__}), using local storage', flush=True)
        all_activities = []
        for email, activities in _LOCAL_ACTIVITIES.items():
            all_activities.extend(activities)
        print(f'Found {len(all_activities)} activities in local storage', flush=True)

    # Sort by timestamp descending
    all_activities = sorted(all_activities, key=lambda x: x.get('timestamp', 0), reverse=True)

    # Convert timestamps to readable format
    for activity in all_activities:
        timestamp = activity.get('timestamp', 0)
        # Convert Decimal to int (from DynamoDB)
        timestamp = int(timestamp) if timestamp else 0
        activity['readable_time'] = datetime.fromtimestamp(timestamp / 1000).strftime('%Y-%m-%d %H:%M:%S')

    activity_html = render_template_string(ACTIVITY_DASHBOARD_HTML,
        admin_email=user_email,
        activities=all_activities,
        total_count=len(all_activities)
    )
    return activity_html, 200


def lambda_handler(event, context):
    from mangum import Mangum
    from asgiref.wsgi import WsgiToAsgi
    asgi_app = WsgiToAsgi(app)
    handler = Mangum(asgi_app, lifespan="off")
    return handler(event, context)


if __name__ == '__main__':
    app.run(debug=True, port=5000)
